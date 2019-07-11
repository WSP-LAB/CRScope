import os
import time
import operator
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlsxwriter

class Report(object):
    def __init__(self, engine, version):
        self.engine = engine
        self.version = version
        self.root_dir = os.path.dirname(os.path.realpath(__file__))
        self.image_dir = '%s/../result/img/%s/%s/' %(self.root_dir, self.engine, time.strftime("%Y%m%d%H%m%s"))
        self.report_dir = '%s/../result/report/%s/%s/' %(self.root_dir, self.engine, time.strftime("%Y%m%d%H%m%s"))

        # create folder
        if os.path.isdir(self.image_dir) == False:
            os.system('mkdir -p %s' %self.image_dir)
        if os.path.isdir(self.report_dir) == False:
            os.system('mkdir -p %s' %self.report_dir)

class Docx(Report):
    def __init__(self, dataset, engine, version):
        super(Docx, self).__init__(engine, version)
        self.doc = Document()
        self.init_doc(dataset)
 
    def close(self):
        self.doc.save(self.report_dir + 'Report.v%s.docx' %self.version)

    def init_doc(self, dataset):
        self.doc.add_paragraph('').add_run('Data Set').bold = True
    
        # write data set count
        count = dataset.count()
        row = 0
        count_table = self._add_table(2, len(count)+1)
        for i, key in enumerate(count.keys()):
            self._add_text(count_table, row, i, key)
            self._add_text(count_table, row+1, i, count[key])

    def write(self, name, figname, info):
        self._write_info(name, info)
        self._write_score(name, figname)

    # write feature size
    def _write_info(self, name, info):
        if len(self.doc.paragraphs) < 3:
            self.doc.add_paragraph('').add_run('\n\nFeatures').bold = True

        feature_table = self._get_table(1)
        if feature_table is None:
            ## create feature table
            row = 0
            feature_table = self._add_table(1, 3)
            self._add_text(feature_table, row, 1, "all features")
            self._add_text(feature_table, row, 2, "selected features")

        ## add feature size
        row = len(feature_table.rows)
        self._add_row(feature_table)
        self._add_text(feature_table, row, 0, name)
        self._add_text(feature_table, row, 1, info[0])

    # add graph image
    def _write_score(self, name, figname):
        sec = self._get_section(1)
        if len(self.doc.paragraphs) < 5:
            self.doc.add_paragraph('').add_run('Evaluation').bold = True

        score_table = self._get_table(2)
        if score_table is None:
            ## create score table
            row = 0
            score_table = self._add_table(1, 3)
            self._add_text(score_table, row, 1, "accuracy")
            self._add_text(score_table, row, 2, "auc")

        ## set column size
        score_table.columns[0].width = Inches(1.2)
        score_table.columns[1].width = Inches(2.4)
        score_table.columns[2].width = Inches(2.4)

        ## add image
        row = len(score_table.rows)
        self._add_row(score_table)
        self._add_text(score_table, row, 0, name)
        self._add_img(score_table, row, 1, (self.image_dir+'accuracy-%s.png' %figname))
        self._add_img(score_table, row, 2, (self.image_dir+'auc-%s.png' %figname))

    def _get_section(self, num):
        if len(self.doc.sections) > num:
            return self.doc.sections[num]
        else:
            return self.doc.add_section()

    def _get_table(self, num):
        if len(self.doc.tables) > num:
            return self.doc.tables[num]
        return None
    
    def _add_table(self, row, col):
        table = self.doc.add_table(rows=row, cols=col)
        table.autofit = True
        return table
    
    def _add_row(self, table):
        table.add_row().cells

    def _add_text(self, table, row, col, text, center=False):
        if center == False:
            table.rows[row].cells[col].text = str(text)
        else:
            p = table.rows[row].cells[col].add_paragraph(str(text))
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_img(self, table, row, col, img):
        p = table.rows[row].cells[col].add_paragraph()
        r = p.add_run()
        r.add_picture(img, width=Inches(2.2))
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
class Xlsx(Report):
    def __init__(self, engine, version):
        super(Xlsx, self).__init__(engine, version)
        self.workbook = xlsxwriter.Workbook(self.report_dir + 'Report_Appendix.v%s.xlsx' %self.version)

        self.header_format = self.workbook.add_format({
                'bold': 1,
                'align': 'center',
                'valign': 'vcenter'
        })
        self.content_format = self.workbook.add_format({
                'align': 'left',
                'valign': 'top'
        })
        
        self.col_list = {}
        self.row_list = {}

    def write(self, name, model):
        self._write_coef(name, model)
        self._write_score(name, model)

    def close(self):
        self.workbook.close()

    def _get_row(self, name):
        if name in self.row_list.keys():
            row = self.row_list[name][0]
        else:
            sorted_row = sorted(self.row_list.items(), key=operator.itemgetter(1))
            row = sorted_row[-1][1][1]+1
            self.row_list[name] = [row, row]
        return row

    def _write_coef(self, name, model):
        worksheet = None

        # find coef workseet
        for ws in self.workbook.worksheets():
            if ws.get_name() == 'coef':
                worksheet = ws
                break
        # create coef worksheet
        if worksheet is None:
            worksheet = self.workbook.add_worksheet('coef')
            self.col_list['coef'] = 1
            self.row_list[name] = [2, 2]
    
        # write column header
        col = self.col_list['coef']
        worksheet.set_column(col, col+1, 21)
        worksheet.merge_range(0, col, 0, col+1, model.name, self.header_format)
        worksheet.write(1, col, 'coef[0]', self.header_format)
        worksheet.write(1, col+1, 'coef[1]', self.header_format)

        # write coef
        row = self._get_row(name)
        cur_row = row
        for coef_list in model.coef:
            for i, coef in enumerate(coef_list):
                worksheet.write(cur_row, col+i, coef, self.content_format)
            worksheet.set_row(cur_row, 70)
            cur_row += 1
        worksheet.merge_range(row, 0, cur_row-1, 0, name, self.header_format)
        
        self.col_list['coef'] += (i+1)
        self.row_list[name] = [row, cur_row-1]

    def _write_score(self, name, model):
        worksheet = None

        # fine score worksheet
        for ws in self.workbook.worksheets():
            if ws.get_name() == 'score':
                worksheet = ws
                break
        # create score worksheet
        if worksheet is None:
            worksheet = self.workbook.add_worksheet('score')
            self.col_list['score'] = 1

        # write column header
        col = self.col_list['score']

        worksheet.set_column(col, col, 9)
        worksheet.set_column(col+1, col+1, 21)
        worksheet.set_column(col+2, col+2, 40)
        worksheet.merge_range(0, col, 0, col+3, model.name, self.header_format)
        worksheet.write(1, col, 'accuracy', self.header_format)
        worksheet.write(1, col+1, 'confusion_matrix', self.header_format)
        worksheet.write(1, col+2, 'classification_report', self.header_format)
        worksheet.write(1, col+3, 'auc', self.header_format)

        # write score
        row = self._get_row(name)
        cur_row = row
        for acc, conf, rep, auc in zip(model.accuracy_score, model.confusion_matrix, model.report, model.roc_auc_score):
            worksheet.write(cur_row, col, acc, self.content_format)
            worksheet.write(cur_row, col+1, '\n'.join(str(conf[i]) for i in range(len(conf))), self.content_format)
            worksheet.write(cur_row, col+2, rep, self.content_format)
            worksheet.write(cur_row, col+3, auc, self.content_format)
            worksheet.set_row(cur_row, 60)
            cur_row += 1
        worksheet.merge_range(row, 0, cur_row-1, 0, name, self.header_format)

        self.col_list['score'] += 4

    def reset_col(self):
        self.col_list['coef'] = 1
        self.col_list['score'] = 1
